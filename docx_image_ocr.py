#!/usr/bin/env python3
import zipfile
import io
import os

from docx import Document
from PIL import Image
import pytesseract

# ——— CONFIGURE THESE ———
# Either hard-code your paths here:
input_docx = 'JamesYoung.docx'
output_docx = 'output.docx'
# Or uncomment to prompt at runtime:
# input_docx = input("Path to input .docx: ").strip()
# output_docx = input("Path for output .docx: ").strip()
# ————————————————————

def extract_images_in_order(docx_path):
    """
    Extract images from a .docx in the visual order they appear, by walking
    through all paragraphs and runs and grabbing any embedded pictures.
    Returns list of (filename, image_bytes).
    """
    doc = Document(docx_path)
    images = []
    # Walk through each paragraph and table cell in sequence
    for block in doc.element.body:
        # newer format drawings
        for drawing in block.findall(".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"):
            rId = drawing.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
            part = doc.part.related_parts[rId]
            filename = os.path.basename(part.partname)
            images.append((filename, part.blob))
        # legacy pictures
        for pict in block.findall(".//{http://schemas.openxmlformats.org/officeDocument/2006/word}binData"):
            rId = pict.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id")
            if rId and rId in doc.part.related_parts:
                part = doc.part.related_parts[rId]
                filename = os.path.basename(part.partname)
                images.append((filename, part.blob))
    return images

def ocr_image_bytes(image_bytes):
    """Run Tesseract OCR on image bytes."""
    with Image.open(io.BytesIO(image_bytes)) as img:
        return pytesseract.image_to_string(img)

def build_output_doc(image_texts, output_path):
    """Create a .docx with headings and OCR text."""
    out = Document()
    out.add_heading('OCR Extracted Text', level=1)
    for idx, (fname, text) in enumerate(image_texts, 1):
        out.add_heading(f'Image {idx}: {fname}', level=2)
        for line in text.splitlines():
            if line.strip():
                out.add_paragraph(line)
    out.save(output_path)

def main():
    if not os.path.isfile(input_docx):
        print(f"Error: input file not found: {input_docx}")
        return

    print("Extracting images…")
    images = extract_images_in_order(input_docx)
    count = len(images)
    if count == 0:
        print("No images found in the document.")
        return

    print(f"Found {count} image{'s' if count != 1 else ''}.")
    image_texts = []
    for fname, img_bytes in images:
        print(f"OCR on {fname}…")
        text = ocr_image_bytes(img_bytes)
        image_texts.append((fname, text))

    print("Building output document…")
    build_output_doc(image_texts, output_docx)
    print(f"DONE. OCR text written to {output_docx}")

if __name__ == '__main__':
    main()
